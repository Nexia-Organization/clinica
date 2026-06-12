"""
Parser para documentos Word de pase de guardia.
Convierte el formato del Word al formato JSON del sistema.
"""
import re
import json
from typing import Dict, List, Any, Optional
from datetime import datetime
import docx


def extract_patient_info(text: str) -> Dict[str, Any]:
    """
    Extrae información del paciente desde el texto de una celda.
    Formato esperado: "sala X cama Y\nnombre apellido edad\ndiagnóstico"
    """
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if not lines:
        return {}
    
    patient = {
        "sala": "",
        "cama": "",
        "nombre": "",
        "apellido": "",
        "edad": "",
        "diagnostico": "",
        "tipo": "control"  # default
    }
    
    # Buscar sala y cama (puede estar en cualquier línea)
    sala_match = re.search(r'sala\s+(\d+)', text.lower())
    if sala_match:
        patient["sala"] = f"Sala {sala_match.group(1)}"
    
    cama_match = re.search(r'cama\s+(\d+)', text.lower())
    if cama_match:
        patient["cama"] = cama_match.group(1)
    
    # Buscar nombre - generalmente después de "sala X cama Y" o en líneas siguientes
    # Buscar patrones como "pepito gomez", "pepita messi", etc.
    name_pattern = re.compile(r'^([a-záéíóúñ]+(?:\s+[a-záéíóúñ]+){0,2})\s+(\d+)\s*(?:años?|meses?|mes)', re.IGNORECASE)
    name_pattern2 = re.compile(r'^([a-záéíóúñ]+(?:\s+[a-záéíóúñ]+){1,2})(?:\s+\d+)?', re.IGNORECASE)
    
    for line in lines:
        # Saltar líneas con "sala" o "cama"
        if 'sala' in line.lower() or 'cama' in line.lower() or 'pase' in line.lower():
            continue
        
        # Intentar match con patrón nombre + edad
        match = name_pattern.match(line)
        if match:
            name_parts = match.group(1).split()
            patient["nombre"] = name_parts[0] if name_parts else ""
            patient["apellido"] = " ".join(name_parts[1:]) if len(name_parts) > 1 else ""
            patient["edad"] = match.group(2)
            break
        
        # Intentar match solo nombre
        match2 = name_pattern2.match(line)
        if match2 and not patient["nombre"]:
            name_parts = match2.group(1).split()
            if len(name_parts) >= 2:
                patient["nombre"] = name_parts[0]
                patient["apellido"] = " ".join(name_parts[1:])
            elif len(name_parts) == 1:
                patient["nombre"] = name_parts[0]
    
    # Buscar edad si no se encontró con el nombre
    if not patient["edad"]:
        edad_match = re.search(r'(\d+)\s*(?:años?|años?|meses?|mes)', text.lower())
        if edad_match:
            patient["edad"] = edad_match.group(1)
    
    # Buscar DNI (número de documento) - generalmente un número de 7-8 dígitos
    dni_match = re.search(r'\b(\d{7,8})\b', text)
    if dni_match:
        # Verificar que no sea parte de la edad o cama
        dni_candidate = dni_match.group(1)
        # Si no es edad ni cama, probablemente es DNI
        if not re.search(rf'{dni_candidate}\s*(?:años?|meses?|mes|cama)', text.lower()):
            patient["dni"] = dni_candidate
    
    # Diagnóstico - generalmente después del nombre/edad, o en líneas que no son nombre
    diagnostic_lines = []
    found_name = False
    for line in lines:
        if 'sala' in line.lower() or 'cama' in line.lower() or 'pase' in line.lower():
            continue
        if patient["nombre"] and patient["nombre"].lower() in line.lower():
            found_name = True
            continue
        if found_name or (not patient["nombre"] and line):
            # Verificar que no sea solo un número (DNI)
            if not re.match(r'^\d+$', line):
                diagnostic_lines.append(line)
    
    if diagnostic_lines:
        patient["diagnostico"] = " ".join(diagnostic_lines)
    
    return patient


def parse_word_document(file_path: str) -> Dict[str, Any]:
    """
    Parsea un documento Word de pase de guardia y lo convierte al formato del sistema.
    """
    doc = docx.Document(file_path)
    
    result = {
        "fecha_guardia": None,
        "ingresos": [],
        "controles": [],
        "pases_ucip": [],
        "pases_a_ucip": []
    }
    
    current_section = None
    current_sala = None
    
    # Procesar párrafos para identificar secciones
    for para in doc.paragraphs:
        text = para.text.strip().upper()
        
        if "GUARDIA" in text:
            # Extraer fecha de guardia
            fecha_match = re.search(r'(\d{1,2})/(\d{1,2})', text)
            if fecha_match:
                result["fecha_guardia"] = text
        
        if "INGRESOS" in text:
            current_section = "ingresos"
            sala_match = re.search(r'SALA\s+(\d+)', text)
            if sala_match:
                current_sala = f"Sala {sala_match.group(1)}"
        
        elif "CONTROLES" in text:
            current_section = "controles"
            sala_match = re.search(r'SALA\s+(\d+)', text)
            if sala_match:
                current_sala = f"Sala {sala_match.group(1)}"
            elif "ONCO" in text:
                current_sala = "Oncología"
        
        elif "PASE" in text and "UCIP" in text:
            if "DE UCIP" in text:
                current_section = "pases_ucip"
            elif "A UCIP" in text:
                current_section = "pases_a_ucip"
    
    # Procesar tablas
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            continue
        
        # Detectar tipo de tabla por encabezados
        header_row = table.rows[0]
        headers = [cell.text.strip().upper() for cell in header_row.cells]
        headers_text = " ".join(headers)
        
        is_ingresos_table = "INDICACIONES" in headers_text
        is_controles_table = "MOTIVO DEL CONTROL" in headers_text
        
        # Determinar sección según contexto y tipo de tabla
        if is_ingresos_table:
            # Tabla de ingresos tiene columna "INDICACIONES"
            section_type = "ingresos"
        elif is_controles_table:
            # Tabla de controles tiene columna "MOTIVO DEL CONTROL"
            section_type = "controles"
        else:
            # Si no se puede determinar por headers, usar la sección actual
            section_type = current_section or "controles"
        
        # Procesar filas de datos (saltar header)
        for row_idx, row in enumerate(table.rows[1:], start=1):
            cells = [cell.text.strip() for cell in row.cells]
            
            if len(cells) < 2:
                continue
            
            # Extraer información del paciente
            patient_info = extract_patient_info(cells[0])
            
            if not patient_info.get("nombre"):
                continue
            
            # Determinar si es pase de UCIP
            if "PASE" in cells[0].upper() and "UCIP" in cells[0].upper():
                if "DE UCIP" in cells[0].upper():
                    section_type = "pases_ucip"
                else:
                    section_type = "pases_a_ucip"
            
            # Construir objeto paciente
            paciente = {
                "patient_id": None,  # Se generará después
                "patient": f"{patient_info.get('nombre', '')} {patient_info.get('apellido', '')}".strip(),
                "nickname": patient_info.get("nombre", ""),
                "nombre": patient_info.get("nombre", ""),
                "apellido": patient_info.get("apellido", ""),
                "sala": patient_info.get("sala") or current_sala or "",
                "cama": patient_info.get("cama", ""),
                "edad": patient_info.get("edad", ""),
                "dni": patient_info.get("dni", ""),
                "diagnostico": patient_info.get("diagnostico", ""),
                "tipo": section_type,
                "indicaciones": cells[1] if len(cells) > 1 else "",
                "examenes": cells[2] if len(cells) > 2 else "",
                "controles": cells[3] if len(cells) > 3 else "",
                "shifts": [],
                "shiftsPharmacy": []
            }
            
            # Generar ID estable basado en nombre + apellido + DNI (si está disponible)
            import hashlib
            nombre = (patient_info.get("nombre", "") or "").lower().strip()
            apellido = (patient_info.get("apellido", "") or "").lower().strip()
            dni = (patient_info.get("dni", "") or "").strip()
            
            # Construir string para ID: nombre + apellido + DNI
            name_for_id = f"{nombre}{apellido}{dni}".replace(" ", "")
            paciente["patient_id"] = f"pat_{hashlib.md5(name_for_id.encode()).hexdigest()[:8]}"
            
            # Agregar a la sección correspondiente
            if section_type == "ingresos":
                result["ingresos"].append(paciente)
            elif section_type == "controles":
                result["controles"].append(paciente)
            elif section_type == "pases_ucip":
                result["pases_ucip"].append(paciente)
            elif section_type == "pases_a_ucip":
                result["pases_a_ucip"].append(paciente)
    
    return result


def convert_to_system_format(parsed_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Convierte los datos parseados al formato del sistema actual.
    """
    all_patients = []
    
    # Combinar todos los pacientes
    for section in ["ingresos", "controles", "pases_ucip", "pases_a_ucip"]:
        for paciente in parsed_data.get(section, []):
            # Crear un shift de enfermería básico con la información
            shift = {
                "shift_id": f"word_import_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                "posted_by": "Sistema de Carga",
                "posted_by_id": "system_001",
                "sent_at": datetime.now().isoformat(timespec="minutes"),
                "shift_label": "Carga desde Word",
                "unit": paciente.get("sala", ""),
                "vitals": {},
                "note": f"Diagnóstico: {paciente.get('diagnostico', 'N/A')}. {paciente.get('controles', '')}",
                "raw_header": f"Importado desde Word - {parsed_data.get('fecha_guardia', 'Fecha no especificada')}",
                "summary_shift": ""
            }
            
            # Crear indicación de farmacia si hay indicaciones
            if paciente.get("indicaciones"):
                pharmacy_shift = {
                    "date": datetime.now().isoformat(timespec="minutes"),
                    "responsableNurse": "Sistema de Carga",
                    "responsable_id": "system_001",
                    "resumen": paciente["indicaciones"]
                }
                paciente["shiftsPharmacy"] = [pharmacy_shift]
            
            paciente["shifts"] = [shift]
            all_patients.append(paciente)
    
    return {
        "patients": all_patients,
        "meta": {
            "source": "word_import",
            "fecha_guardia": parsed_data.get("fecha_guardia"),
            "total_pacientes": len(all_patients),
            "ingresos": len(parsed_data.get("ingresos", [])),
            "controles": len(parsed_data.get("controles", [])),
            "pases_ucip": len(parsed_data.get("pases_ucip", [])),
            "pases_a_ucip": len(parsed_data.get("pases_a_ucip", []))
        }
    }


if __name__ == "__main__":
    # Test
    parsed = parse_word_document("pacientes.docx")
    print(json.dumps(parsed, indent=2, ensure_ascii=False))
    
    converted = convert_to_system_format(parsed)
    print("\n=== CONVERTIDO ===")
    print(json.dumps(converted, indent=2, ensure_ascii=False)[:2000])
